#!/usr/bin/env python3
"""
generate_surat_pernyataan.py
---------------------------
Membuat dokumen SURAT PERNYATAAN TANGGUNG JAWAB MUTLAK (SPTJM)
untuk pembayaran TPP dari data template yang ada di sheet 'sptjm'
dalam file Excel TPP.

Output dapat berupa Word (.docx) maupun HTML (.html) yang
mengikuti struktur file html.html.

Cara pakai:
    # Word (docx) - default
    python generate_surat_pernyataan.py "TPP PNS JULI 2026 SMKN 1 KOBA.xlsm" -t PNS \
        -o "SURAT_PERNYATAAN_PNS.docx"

    python generate_surat_pernyataan.py "TPP P3K JULI 2026 SMKN 1 KOBA.xlsm" -t PPPK \
        -o "SURAT_PERNYATAAN_P3K.docx"

    # HTML
    python generate_surat_pernyataan.py "TPP PNS JULI 2026 SMKN 1 KOBA.xlsm" -t PNS -f html \
        -o "SURAT_PERNYATAAN_PNS.html"

Output:
    File Word (.docx) atau HTML (.html) berisi surat pernyataan dengan format standir
"""

import argparse
import base64
import re
import sys
from html import escape
from pathlib import Path

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def extract_sptjm_data(xlsx_path):
    """Baca sheet 'sptjm' dan kembalikan dictionary data surat."""
    wb = load_workbook(xlsx_path, data_only=True)
    if "sptjm" not in wb.sheetnames:
        raise ValueError(f"Sheet 'sptjm' tidak ditemukan di {xlsx_path}. "
                         f"Sheet tersedia: {wb.sheetnames}")

    ws = wb["sptjm"]

    data = {
        "nama": "",
        "nip": "",
        "jabatan": "",
        "spm_nomor": "",
        "spm_tanggal": "",
        "jumlah_rupiah": "",
        "jumlah_terbilang": "",
        "bulan": "",
        "tahun": "",
        "lokasi": "",
        "bulan_tahun_tempat": "",
    }

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True):
        for idx, cell in enumerate(row):
            val = str(cell).strip() if cell is not None else ""
            if not val:
                continue

            if "Nama" in val and idx + 1 < len(row) and row[idx + 1]:
                data["nama"] = str(row[idx + 1]).strip()

            if val == "NIP" and idx + 1 < len(row) and row[idx + 1]:
                data["nip"] = str(row[idx + 1]).strip()

            if "Jabatan" in val and idx + 1 < len(row) and row[idx + 1]:
                data["jabatan"] = str(row[idx + 1]).strip()

            if "Nomor" in val and ":" in val:
                after = val.split(":", 1)[1].strip()
                dots = re.findall(r"\.{3,}", after)
                if dots:
                    data["spm_nomor"] = after.split(dots[0])[0].strip()

            if "tanggal" in val.lower():
                parts = re.split(r"tanggal\s*", val, flags=re.IGNORECASE)
                if len(parts) > 1:
                    after = parts[1].strip()
                    dots = re.findall(r"\.{3,}", after)
                    if dots:
                        data["spm_tanggal"] = after.split(dots[0])[0].strip()

            if re.search(r"Rp\.?\s*[\d\.]+,-", val):
                m = re.search(r"(Rp\.?\s*[\d\.]+,-)", val)
                if m:
                    data["jumlah_rupiah"] = m.group(1).strip()
                m2 = re.search(r"Rp\.?\s*[\d\.]+,\-\s*\(([^)]+)\)", val)
                if m2:
                    data["jumlah_terbilang"] = m2.group(1).strip()

            if re.search(r"untuk bulan\s+\w+\s+\d{4}", val, re.IGNORECASE):
                m = re.search(r"untuk bulan\s+(\w+)\s+(\d{4})", val, re.IGNORECASE)
                if m:
                    data["bulan"] = m.group(1).strip()
                    data["tahun"] = m.group(2).strip()

            if re.match(r"^Koba,\s*\w+\s+\d{4}$", val.strip()):
                data["lokasi"] = val.strip()
                parts = val.split(",")
                if len(parts) >= 2:
                    data["bulan_tahun_tempat"] = parts[1].strip()

    return data


def set_paragraph_font(paragraph, font_name="Arial", font_size=11, bold=False):
    """Terapkan font ke semua run di paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT,
                  font_name="Arial", font_size=11, bold=False,
                  space_after=6, space_before=0):
    """Tambah paragraph dengan gaya yang konsisten."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    return p


def add_tab_paragraph(doc, label, value, tab_pos_cm=4.5,
                       font_name="Arial", font_size=11, bold=False,
                       space_after=3, space_before=0):
    """Tambah paragraph dengan label di-sepakati kanan, colon, lalu value.
    Contoh output:
        Nama       : SYAHRYANTO, S.T.,M.Pd
        NIP        : 197708262006041005
        Jabatan    : Kepala SMK Negeri 1 Koba
    sehingga titik dua (:) align vertikal.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_pos_cm))

    run_label = p.add_run(label)
    run_label.font.name = font_name
    run_label.font.size = Pt(font_size)
    run_label.font.bold = bold

    run_tab = p.add_run("\t")
    run_tab.font.name = font_name
    run_tab.font.size = Pt(font_size)

    run_colon = p.add_run(": ")
    run_colon.font.name = font_name
    run_colon.font.size = Pt(font_size)
    run_colon.font.bold = bold

    run_value = p.add_run(value)
    run_value.font.name = font_name
    run_value.font.size = Pt(font_size)
    run_value.font.bold = bold

    return p


def add_numbered_paragraph(doc, text, font_name="Arial", font_size=11,
                           space_after=12, space_before=0):
    """Tambah numbered paragraph (ordered list) untuk item 1, 2, 3, dst.
    Gunakan style 'List Number' agar Word menangani penomoran otomatis.
    Alignment LEFT agar baris terakhir tidak terlalu renggang.
    """
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.left_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.add_break()

    return p


def add_indented_paragraph(doc, text, indent_px=200,
                           font_name="Arial", font_size=11, bold=False,
                           align=WD_ALIGN_PARAGRAPH.LEFT,
                           space_after=6, space_before=0):
    """Tambah paragraph dengan left indent dalam pixel."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.left_indent = Emu(indent_px * 12700)

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    return p


def get_surat_texts(data, surat_type="PNS"):
    """Kembalikan 3 string isi bernomor surat ( dipakai builder docx & html)."""
    nomor_spm = data.get("spm_nomor") or "..........................................."
    tanggal_spm = data.get("spm_tanggal") or ".........................................."
    jumlah = data.get("jumlah_rupiah") or ""
    terbilang = data.get("jumlah_terbilang") or ""
    bulan = data.get("bulan") or ""
    tahun = data.get("tahun") or ""

    if not jumlah and not terbilang:
        jumlah = "Rp. ....................................,-"
        terbilang = "..........................................."

    pegawai_type = "Negeri Sipil" if surat_type == "PNS" else "PPPK"

    item1 = (
        f"Perhitungan yang terdapat dalam SPM Langsung (SPM-LS) Nomor : {nomor_spm} "
        f"tanggal {tanggal_spm} untuk pembayaran Tambahan Penghasilan Pegawai (TPP) "
        f"{pegawai_type} sebesar {jumlah} ({terbilang}) untuk bulan {bulan} {tahun} "
        f"telah dihitung dengan benar berdasarkan dokumen pelaksanaan anggaran "
        f"dan dokumen pendukung lainnya."
    )
    item2 = (
        "Apabila terdapat kesalahan dan kelebihan atas pembayaran, sebagaimana "
        "yang dimaksud pada point 1 (satu), kami bertanggung jawab dan bersedia "
        "untuk menyetorkan kelebihan tersebut ke Kas Daerah."
    )
    item3 = (
        "Dokumen bukti-bukti belanja atas pembayaran tersebut di atas disimpan di "
        "Dinas Pendidikan Provinsi Kepulauan Bangka Belitung (SMK Negeri 1 Koba) "
        "sesuai ketentuan yang berlaku untuk kelengkapan administrasi dan keperluan "
        "pemeriksaan BPK dan/atau aparatur pengawas fungsional lainnya."
    )
    return item1, item2, item3


SURAT_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <title>Surat Pernyataan Tanggung Jawab Mutlak (SPTJM) @@TIPE@@</title>
    <style>
        @page {
            size: A4;
            margin: 10mm 20mm 20mm 20mm;
            background-color: #ffffff;
        }
        *, *::before, *::after {
            box-sizing: border-box;
        }
        body {
            font-family: 'arial', sans-serif;
            font-size: 12pt;
            line-height: 1.5;
            color: #000000;
            margin: 0;
            padding: 0;
        }
        .kop-image {
            max-width: 560px;
            width: 100%;
            height: auto;
            display: block;
            margin: 0 auto 10px;
        }
        .title-container {
            text-align: center;
            margin-top: 20px;
            margin-bottom: 25px;
        }
        .title-container h4 {
            margin: 0;
            font-size: 12pt;
            font-weight: bold;
            text-decoration: underline;
            text-transform: uppercase;
        }
        .content-section {
            margin-bottom: 15px;
            text-align: justify;
        }
        .identity-table {
            width: 100%;
            margin-bottom: 15px;
            border-collapse: collapse;
            margin-left: 30px;
        }
        .identity-table td {
            vertical-align: top;
            padding: 2px 0;
        }
        .identity-table td.label {
            width: 10%;
        }
        .identity-table td.separator {
            width: 3%;
            text-align: center;
        }
        .identity-table td.value {
            width: 75%;
        }
        .numbered-list {
            margin: 0;
            padding-left: 20px;
            text-align: justify;
        }
        .numbered-list li {
            margin-bottom: 10px;
            padding-left: 5px;
        }
        .signature-section {
            margin-top: 40px;
            float: right;
            width: 280px;
            text-align: left;
        }
        .signature-section p {
            margin: 0 0 5px 0;
        }
        .signature-space {
            height: 70px;
        }
        .signatory-name {
            font-weight: normal;
            /* text-decoration: underline; */
        }
    </style>
</head>
<body>
@@HEADER@@
<div class="title-container">
    <h4>SURAT PERNYATAAN TANGGUNG JAWAB MUTLAK</h4>
</div>
<div class="content-section">
    <p>Yang bertanda tangan dibawah ini :</p>
    <table class="identity-table">
        <tr>
            <td class="label">Nama</td>
            <td class="separator">:</td>
            <td class="value"><b>@@NAMA@@</b></td>
        </tr>
        <tr>
            <td class="label">NIP</td>
            <td class="separator">:</td>
            <td class="value">@@NIP@@</td>
        </tr>
        <tr>
            <td class="label">Jabatan</td>
            <td class="separator">:</td>
            <td class="value">@@JABATAN@@</td>
        </tr>
    </table>
    <p>Menyatakan dengan sesungguhnya bahwa :</p>
</div>
<div class="content-section">
    <ol class="numbered-list">
        <li>@@ITEM1@@</li>
        <li>@@ITEM2@@</li>
        <li>@@ITEM3@@</li>
    </ol>
</div>
<div class="content-section">
    <p>Demikian surat pernyataan ini dibuat dengan sesungguhnya untuk dipergunakan sebagaimana mestinya.</p>
</div>
<div class="signature-section">
    <p>@@LOKASI@@</p>
    <p>Kepala Sekolah,</p>
    <div class="signature-space"></div>
    <p class="signatory-name">@@NAMA@@<br>NIP. @@NIP@@</p>
</div>
</body>
</html>
"""


def build_surat_html(data, output_path, surat_type="PNS"):
    """Bangun file HTML surat pernyataan berdasarkan struktur html.html."""
    item1, item2, item3 = get_surat_texts(data, surat_type)
    lokasi_tanggal = data.get("lokasi") or "Koba, Agustus 2026"

    kop_path = Path(__file__).parent / "kop_surat.jpg"
    if kop_path.exists():
        b64 = base64.b64encode(kop_path.read_bytes()).decode("ascii")
        header_html = f'<img src="data:image/jpeg;base64,{b64}" class="kop-image" alt="Kop Surat">'
    else:
        header_html = ""

    html = SURAT_HTML_TEMPLATE
    html = html.replace("@@TIPE@@", escape(surat_type))
    html = html.replace("@@HEADER@@", header_html)
    html = html.replace("@@NAMA@@", escape(data.get("nama") or ""))
    html = html.replace("@@NIP@@", escape(data.get("nip") or ""))
    html = html.replace("@@JABATAN@@", escape(data.get("jabatan") or ""))
    html = html.replace("@@ITEM1@@", escape(item1))
    html = html.replace("@@ITEM2@@", escape(item2))
    html = html.replace("@@ITEM3@@", escape(item3))
    lokasi_tanggal = lokasi_tanggal.replace(",", ",<span style='margin-right:30px;'>&nbsp;</span>")
    html = html.replace("@@LOKASI@@", lokasi_tanggal)

    output_path = Path(output_path)
    output_path.write_text(html, encoding="utf-8")
    print(f"Surat pernyataan HTML disimpan ke: {output_path}")
    return output_path


def build_surat(data, output_path, surat_type="PNS"):
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    kop_path = Path(__file__).parent / "kop_surat.jpg"
    if kop_path.exists():
        doc.add_picture(str(kop_path), width=Inches(5.8))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.paragraph_format.space_after = Pt(6)

    add_paragraph(doc, "SURAT PERNYATAAN TANGGUNG JAWAB MUTLAK",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  font_size=14, space_after=12)

    add_paragraph(doc, "Yang bertanda tangan dibawah ini :", space_after=6)

    add_tab_paragraph(doc, "Nama", data['nama'], space_after=3)
    add_tab_paragraph(doc, "NIP", data['nip'], space_after=3)
    add_tab_paragraph(doc, "Jabatan", data['jabatan'], space_after=12)

    add_paragraph(doc, "Menyatakan dengan sesungguhnya bahwa :", space_after=6)

    item1, item2, item3 = get_surat_texts(data, surat_type)
    add_numbered_paragraph(doc, item1, space_after=12)
    add_numbered_paragraph(doc, item2, space_after=12)
    add_numbered_paragraph(doc, item3, space_after=12)

    lokasi_tanggal = data.get("lokasi") or "Koba, Agustus 2026"
    if not lokasi_tanggal.startswith("Koba"):
        lokasi_tanggal = f"Koba, {lokasi_tanggal}"

    add_indented_paragraph(doc, lokasi_tanggal, indent_px=200, space_after=6)
    add_indented_paragraph(doc, "Kepala Sekolah", indent_px=200, space_after=36)

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_after = Pt(3)
    p_sig.paragraph_format.space_before = Pt(0)
    p_sig.paragraph_format.left_indent = Emu(200 * 12700)
    run_name = p_sig.add_run(data["nama"])
    run_name.font.name = "Arial"
    run_name.font.size = Pt(11)
    run_name.font.bold = True
    run_name.add_break()
    run_nip = p_sig.add_run(f"NIP. {data['nip']}")
    run_nip.font.name = "Arial"
    run_nip.font.size = Pt(11)

    output_path = Path(output_path)
    doc.save(str(output_path))
    print(f"Surat pernyataan disimpan ke: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Generate SURAT_PERNYATAAN_PNS / SURAT_PERNYATAAN_P3K (.docx atau .html) dari sheet sptjm Excel TPP")
    parser.add_argument("xlsx_path", help="Path ke file Excel TPP (mis. TPP PNS JULI 2026 SMKN 1 KOBA.xlsm)")
    parser.add_argument("-t", "--type", default="PNS", choices=["PNS", "PPPK"],
                        help="Tipe surat: PNS atau PPPK (default: PNS)")
    parser.add_argument("-f", "--format", default="docx", choices=["docx", "html"],
                        help="Format output: docx atau html (default: docx)")
    parser.add_argument("-o", "--output", default=None,
                        help="Path output (default: SURAT_PERNYATAAN_PNS.docx/.html atau SURAT_PERNYATAAN_P3K.docx/.html)")
    args = parser.parse_args()

    xlsx_path = Path(args.xlsx_path)
    if not xlsx_path.exists():
        sys.exit(f"File tidak ditemukan: {xlsx_path}")

    ext = "html" if args.format == "html" else "docx"
    if args.output:
        output_path = Path(args.output)
    else:
        if args.type == "PPPK":
            output_path = xlsx_path.with_name(f"SURAT_PERNYATAAN_P3K.{ext}")
        else:
            output_path = xlsx_path.with_name(f"SURAT_PERNYATAAN_PNS.{ext}")

    data = extract_sptjm_data(xlsx_path)
    if args.format == "html":
        build_surat_html(data, output_path, surat_type=args.type)
    else:
        build_surat(data, output_path, surat_type=args.type)

    print(f"  Nama   : {data['nama']}")
    print(f"  NIP    : {data['nip']}")
    print(f"  Jabatan: {data['jabatan']}")
    print(f"  Bulan  : {data['bulan']} {data['tahun']}")
    print(f"  Jumlah : {data['jumlah_rupiah']}")


if __name__ == "__main__":
    main()
